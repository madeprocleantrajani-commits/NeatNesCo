"""
Generate Neatnestco Business Plan DOCX — based on live bot intelligence data.
"""
import json
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
    return h


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    return table


def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_kpi_box(doc, label, value, note=""):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(str(value))
    run2.font.size = Pt(14)
    run2.bold = True
    run2.font.color.rgb = RGBColor(0x00, 0xB8, 0x94)
    if note:
        run3 = p.add_run(f"  ({note})")
        run3.font.size = Pt(9)
        run3.font.color.rgb = RGBColor(0x63, 0x6E, 0x72)


# ── Load intelligence data ───────────────────────────────────

data_files = sorted(Path("data").glob("analysis_*.json"), reverse=True)
analysis = {}
if data_files:
    with open(data_files[0]) as f:
        analysis = json.load(f)

sources = analysis.get("sources", {})
amazon = analysis.get("amazon", {})
candidates = amazon.get("dropship_candidates", [])
cross_refs = analysis.get("cross_references", [])
niches = analysis.get("niche_scores", {})
trends = analysis.get("trends", {})
ali = analysis.get("aliexpress", {})
categories = amazon.get("categories", {})
dist = amazon.get("price_distribution", {})
rising = amazon.get("rising_products", [])
movers = amazon.get("movers", [])
today = datetime.now().strftime("%Y-%m-%d")

active_sources = sum(1 for s in sources.values() if s.get("active"))
total_sources = len(sources)


# ── Create Document ──────────────────────────────────────────

doc = Document()

# Set default font
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x2D, 0x34, 0x36)

# ═══════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("NEATNESTCO")
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PLANI I BIZNESIT — BUSINESS PLAN")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x74, 0xB9, 0xFF)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Dropshipping Intelligence Platform")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x63, 0x6E, 0x72)

doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Data: {today}")
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Powered by Signal Pilot v4 — AI Intelligence Engine")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x63, 0x6E, 0x72)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("8hendy-8i.myshopify.com")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x00, 0xB8, 0x94)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "TABELA E PERMBAJTJES / TABLE OF CONTENTS", 1)
toc_items = [
    "1. Permbledhje Ekzekutive / Executive Summary",
    "2. Avantazhi Teknologjik / Technology Advantage",
    "3. Analiza e Tregut / Market Analysis (Live Data)",
    "4. Strategjia e Produkteve / Product Strategy",
    "5. Nishet me Potencial / Niche Opportunities",
    "6. Strategjia e Marketingut / Marketing Strategy",
    "7. Modeli Financiar / Financial Model",
    "8. Plani Operacional / Operations Plan",
    "9. Plani i Veprimit 30-60-90 Dite / Action Plan",
    "10. Risqet dhe Zbutja / Risks & Mitigation",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "1. PERMBLEDHJE EKZEKUTIVE / EXECUTIVE SUMMARY", 1)

doc.add_paragraph(
    "Neatnestco eshte nje biznes dropshipping i bazuar ne inteligjence artificiale "
    "qe perdor nje sistem te avancuar botesh per te identifikuar produkte fitimprurese "
    "ne tregun amerikan. Platforma operon 24/7 nga nje VPS ne Virginia (us-east-1) "
    "dhe monitoron 7 burime te dhenash ne kohe reale."
)

doc.add_paragraph(
    "Neatnestco is an AI-powered dropshipping business that uses an advanced bot system "
    "to identify profitable products in the US market. The platform operates 24/7 from "
    "a Virginia VPS server, monitoring 7 real-time data sources simultaneously."
)

add_heading_styled(doc, "Shifra Kyqe / Key Metrics (Live)", 2)
add_kpi_box(doc, "Produkte te Skanuara", f"{dist.get('total', 0)}", "Amazon BSR")
add_kpi_box(doc, "Kandidate Dropship", f"{len(candidates)}", "$15-$80, 4.0+ rating, non-brand")
add_kpi_box(doc, "Multi-Signal Matches", f"{len(cross_refs)}", "2+ burime konfirmojne")
add_kpi_box(doc, "Nishe te Analizuara", f"{len(niches)}", "10 nishe me scoring")
add_kpi_box(doc, "Burime Aktive", f"{active_sources}/{total_sources}", "botesh ne pune")
add_kpi_box(doc, "Kategori Amazon", f"{len(categories)}", "monitoruar ne kohe reale")

doc.add_paragraph("")
doc.add_paragraph(
    "MISIONI: Te ndertojme nje biznes dropshipping fitimprures duke perdorur inteligjencen "
    "e te dhenave per te zgjedhur produktet me marzh te larte, kerkese te vertetuar, "
    "dhe konkurrence te ulet."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. TECHNOLOGY ADVANTAGE
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "2. AVANTAZHI TEKNOLOGJIK / TECHNOLOGY ADVANTAGE", 1)

doc.add_paragraph(
    "Neatnestco ka nje avantazh te madh konkurrues: nje sistem inteligjence i automatizuar "
    "qe shumica e dropshipper-ave nuk e kane. Kjo na jep te dhena ne kohe reale per vendimmarrje."
)

add_heading_styled(doc, "Arkitektura e Sistemit", 2)

bot_table = [
    ("Trend Scanner", "Google Trends", "Identifikon produkte ne rritje me regresion linear"),
    ("Amazon Tracker", "Amazon BSR", f"Skanon {dist.get('total', 0)} produkte, {len(movers)} movers, {len(rising)} rising"),
    ("AliExpress Scanner", "AliExpress", "Gjen furnizues me cmim te ulet, kalkulon fitim real"),
    ("eBay Scanner", "eBay Sold", "Validon kerkesen reale me shitje te vertetuara"),
    ("Price Monitor", "Multiple", "Gjurmon ndryshimet e cmimeve, alerta per renie"),
    ("Competitor Tracker", "Shopify stores", "Zbulon konkurrentet, analizon strategjite"),
    ("Intelligence Engine", "All Sources", "Scoring multi-sinjal, identifikon fitimet reale"),
    ("AI Analyzer", "Claude AI", "Analiza e thelle me AI per cdo produkt kandidat"),
    ("Report Generator", "All Data", "Raporte HTML bilinguale (EN+SQ) ne Telegram"),
    ("Alert Bot", "Telegram", "Njoftime ne kohe reale per mundesi te reja"),
]

add_table(doc, ["Boti", "Burimi", "Funksioni"], bot_table)

doc.add_paragraph("")
doc.add_paragraph(
    "FORMULA E FITIMIT REAL: Fitimi = Cmimi_Retail - Cmimi_Burim - Transport "
    "- Platforma(15%) - Reklama(30%). Vetem produktet me marzh pozitiv kalojne filtrin."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. MARKET ANALYSIS (LIVE DATA)
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "3. ANALIZA E TREGUT / MARKET ANALYSIS", 1)

doc.add_paragraph(
    "Te dhenat me poshte jane te gjeneruara LIVE nga sistemi i boteve. "
    "Ato perfaqesojne gjendjen aktuale te tregut amerikan."
)

# Price Distribution
add_heading_styled(doc, "Shperndarja e Cmimeve ne Amazon (USD)", 2)
if dist:
    total = dist.get("total", 1)
    price_rows = []
    for label, key in [("Nen $10", "under_10"), ("$10-$25", "10_to_25"),
                       ("$25-$50", "25_to_50"), ("$50-$100", "50_to_100"),
                       ("Mbi $100", "over_100")]:
        count = dist.get(key, 0)
        pct = round(count / total * 100) if total else 0
        sweet = "DA" if label in ("$10-$25", "$25-$50", "$50-$100") else ""
        price_rows.append((label, str(count), f"{pct}%", sweet))

    add_table(doc, ["Gama", "Produkte", "%", "Sweet Spot?"], price_rows)
    doc.add_paragraph(
        f"Mesatare: ${dist.get('avg', 0):.2f}  |  Mediane: ${dist.get('median', 0):.2f}  |  "
        f"Zona e arte: $15-$80 (optimal per dropship)"
    )

# Categories
add_heading_styled(doc, "Kategorite e Amazon-it", 2)
cat_rows = []
for cn, cd in sorted(categories.items(), key=lambda x: x[1].get("total", 0), reverse=True):
    total = cd.get("total", 0)
    avg = cd.get("avg_price", 0)
    rating = cd.get("avg_rating", 0)
    if total > 0:
        cat_rows.append((cn.replace("-", " & ").replace("_", " ").title(),
                        str(total), f"${avg:.2f}", f"{rating:.1f}/5"))
add_table(doc, ["Kategoria", "Produkte", "Avg Cmimi", "Rating"], cat_rows)

# Top Candidates
add_heading_styled(doc, "Top 20 Kandidate Dropship", 2)
doc.add_paragraph("Filtrat: $15-$80 | Rating 4.0+ | Jo brand | I pershtatshem per dropship")
cand_rows = []
for c in candidates[:20]:
    score = c.get("dropship_score", 0)
    price = c.get("price_usd", 0)
    rating = c.get("rating", 0)
    rev = c.get("review_count", 0)
    title = c.get("title", "?")[:40]
    cand_rows.append((str(score), title, f"${price:.2f}", f"{rating:.1f}", f"{rev:,}"))

add_table(doc, ["Score", "Produkti", "Cmimi", "Rating", "Reviews"], cand_rows)

# Cross references
add_heading_styled(doc, f"Multi-Signal Products ({len(cross_refs)} gjetur)", 2)
doc.add_paragraph(
    "Keto produkte u gjeten ne 2+ burime te pavarura, duke konfirmuar kerkesen reale. "
    "Ato jane kandidatet me te forte per listing ne dyqan."
)
ref_rows = []
for ref in cross_refs[:15]:
    signals = ", ".join(ref.get("signals", []))
    title = ref.get("title", "?")[:40]
    sc = ref.get("signal_count", 1)
    ref_rows.append((str(sc), title, signals))
if ref_rows:
    add_table(doc, ["Sinjale", "Produkti", "Burimet"], ref_rows)

# Trends
hot_products = trends.get("hot_products", [])
if hot_products:
    add_heading_styled(doc, "Trende ne Rritje (Google Trends)", 2)
    trend_rows = []
    for p in hot_products[:10]:
        interest = p.get("current_interest", 0)
        trend = p.get("trend", "?")
        kw = p["keyword"]
        trend_rows.append((kw, f"{interest}/100", trend))
    add_table(doc, ["Keyword", "Interesi", "Drejtimi"], trend_rows)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. PRODUCT STRATEGY
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "4. STRATEGJIA E PRODUKTEVE / PRODUCT STRATEGY", 1)

add_heading_styled(doc, "Kriteret e Perzgjedhjes", 2)
criteria = [
    ("Cmimi Retail: ", "$15 - $80 (zona e arte per dropship)"),
    ("Rating: ", "4.0+ yll (besueshmeri e larte)"),
    ("Reviews: ", "100-10,000 (treg i validuar por jo i ngopir)"),
    ("Brand: ", "Jo brand te medha (shmangim i problemeve te licensimit)"),
    ("Marzhi: ", "20%+ pas te gjitha kostove (burim+transport+platforme+reklama)"),
    ("Kerkesa: ", "Konfirmuar nga 2+ burime (Amazon+eBay+Trends)"),
    ("Konkurrenca: ", "E ulet deri mesatare (mundesi per pozicionim)"),
]
for bold, text in criteria:
    add_bullet(doc, text, bold)

add_heading_styled(doc, "Kategori Prioritare (sipas te dhenave)", 2)
doc.add_paragraph(
    "Bazuar ne analizen e boteve, keto jane kategorite me potencial te larte:"
)

priority_cats = [
    ("Home & Kitchen", "$23.64 avg", "30 produkte aktive", "Produkte shtepie, organizim, gadgets kuzhine"),
    ("Sports & Fitness", "$22.43 avg", "30 produkte aktive", "Pajisje fitnesi, aksesore ushtrimesh, shishka uji"),
    ("Beauty & Health", "$16.61 avg", "30 produkte aktive", "Kujdesi per lekuren, serum flokesh, maska"),
    ("Baby Products", "$25.70 avg", "30 produkte aktive", "Peceta, kujdesi per beben, aksesore"),
    ("Tools & Hardware", "$30.16 avg", "28 produkte aktive", "Vegla dore, filtrat, aksesore shtepie"),
    ("Automotive", "$19.97 avg", "29 produkte aktive", "Aksesore makine, freskonjese, mbrojtje"),
]
add_table(doc, ["Kategoria", "Cmimi Avg", "Produkte", "Pershkrim"], priority_cats)

add_heading_styled(doc, "Produkte te Sugjeruara per Fillim", 2)
doc.add_paragraph(
    "Bazuar ne kryqezimin e te dhenave, keto jane produktet me potencialin me te larte:"
)

# Best product picks based on the data
picks = [
    ("Shower Head (High Pressure)", "$47.46", "7,313", "4.5",
     "Kerkese e larte, marzh i mire, i lehte per transport, tendence ne rritje"),
    ("Under Cabinet LED Lighting", "$15.99", "51,473", "4.5",
     "Cmim i ulet burimi, kerkese masive, i lehte, nuk prishet lehte"),
    ("Water Bottle (Insulated)", "$34.99", "18,194", "4.5",
     "Tendence e forte, multi-signal, treg i madh"),
    ("Weighted Vest (Fitness)", "$28.98", "14,648", "4.5",
     "Nishe fitness ne rritje, marzh i mire, bleresit aktivie shpenzojne"),
    ("Slim Toaster (Kitchen)", "$24.98", "12,378", "4.3",
     "Kategori shtepie, dizajn modern, kerkese e larte"),
    ("Car Air Freshener (Wood)", "$8-15", "1,000+", "4.5+",
     "Cmim i ulet burimi $1-3, marzh i larte, multi-signal ne bote"),
    ("Glass Olive Oil Sprayer", "$12-20", "5,000+", "4.6+",
     "Tendence ne rritje, 2 brande te ndryshme ne top sellers"),
    ("Resistance Bands Set", "$10-20", "10,000+", "4.5+",
     "Nishe fitness, cmim i ulet burimi, marzh i shkelqyer"),
    ("Mattress Protector", "$25.49", "15,806", "4.6",
     "Kerkese konstante, jo sezonale, marzh i mire"),
    ("Air Filter (MERV 8/5)", "$29-40", "100,000+", "4.7+",
     "Kerkese e jashtezakonshme, bleresit riblejne, LTV e larte"),
]
add_table(doc, ["Produkti", "Cmimi", "Reviews", "Rating", "Pse?"], picks)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. NICHE OPPORTUNITIES
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "5. NISHET ME POTENCIAL / NICHE OPPORTUNITIES", 1)

doc.add_paragraph(
    "Sistemi i inteligjences vlereson cdo nishe ne 4 dimensione: "
    "Trend (rritja ne Google), Kerkese (Amazon autocomplete), "
    "Burim (disponueshmeri ne AliExpress), Validim (shitje ne eBay)."
)

add_heading_styled(doc, "Matrica e Nishave (Live Scoring)", 2)

niche_rows = []
for ns in sorted(niches.values(), key=lambda x: x["total"], reverse=True):
    name = ns["name"].replace("_", " ").title()
    total = ns["total"]
    grade = ns["grade"]
    trend = ns.get("trend_score", 0)
    demand = ns.get("demand_score", 0)
    source = ns.get("source_score", 0)
    valid = ns.get("validation_score", 0)
    niche_rows.append((name, str(trend), str(demand), str(source), str(valid),
                       f"{total}/100", grade))

add_table(doc, ["Nisha", "Trend", "Kerkese", "Burim", "Validim", "Total", "Nota"], niche_rows)

doc.add_paragraph("")
doc.add_paragraph(
    "SHENIM: Notat jane te uleta sepse vetem 3/7 burime jane aktive. "
    "Pasi te aktivizohen eBay, Demand, Competitors, dhe Prices, "
    "notat do te rriten ndjeshm. Outdoor dhe Fitness kane potencialin me te madh."
)

add_heading_styled(doc, "Strategjia per Nishe", 2)
add_bullet(doc, "Fillo me 2-3 nishe: Home & Kitchen + Fitness + Beauty", "FAZA 1: ")
add_bullet(doc, "Shto nishe te reja kur marzhi konfirmohet: Pet + Automotive", "FAZA 2: ")
add_bullet(doc, "Ekspando ne nishe sezonale: Outdoor (pranvere/vere), Holiday (dimmer)", "FAZA 3: ")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. MARKETING STRATEGY
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "6. STRATEGJIA E MARKETINGUT / MARKETING STRATEGY", 1)

doc.add_paragraph(
    "Strategjia e marketingut eshte AGRESIVE dhe e bazuar ne te dhena. "
    "Cdo dollar i shpenzuar ne marketing duhet te kete ROI te matshem."
)

add_heading_styled(doc, "Kanalet e Marketingut (sipas prioritetit)", 2)

marketing_channels = [
    ("1. Facebook/Meta Ads", "PRIMAR", "$5-15/dite fillim",
     "Targetim i sakte demografik, lookalike audiences, retargeting. "
     "Fillo me $5/dite per produkt, shkalo ne $50-100 kur ROAS > 2x"),
    ("2. TikTok Ads", "PRIMAR", "$10-20/dite",
     "Video te shkurtra produkti, UGC style. Gen Z dhe Millennials. "
     "CPC me i ulet se Facebook, reach organik i larte"),
    ("3. Google Shopping", "SEKONDAR", "$10-20/dite",
     "Bleresit me intent te larte. Konversion rate me i larte. "
     "Product Listing Ads (PLA) direkt ne kerkim"),
    ("4. Instagram Reels", "SEKONDAR", "Organik + $5/dite",
     "Showcase vizual i produkteve. Story Ads per retargeting. "
     "Influencer micro-partnerships ($50-200 per post)"),
    ("5. Pinterest", "TERCIAR", "Organik",
     "Pins te produkteve. Treg kryesisht femer. Ideal per home, beauty, kitchen. "
     "Traffic evergreen qe rritet me kohen"),
    ("6. SEO/Content", "AFATGJATE", "Organik",
     "Blog ne Shopify: reviews, guides, comparisons. "
     "Keywords long-tail nga Amazon Demand bot. Traffic falas i perjeteshem"),
]
add_table(doc, ["Kanali", "Prioriteti", "Buxheti", "Strategjia"], marketing_channels)

add_heading_styled(doc, "Strategjia e Cmimeve", 2)
doc.add_paragraph("Modeli i cmimeve bazohet ne formulen e fitimit real te botit:")
add_bullet(doc, "Retail = Cmimi_Burim x 2.5 - 3.5 (markup)", "Markup: ")
add_bullet(doc, "Target marzh: 25-40% pas te gjitha kostove", "Marzhi: ")
add_bullet(doc, "Transporti: Free shipping mbi $25 (rrit konversionin 30%+)", "Transporti: ")
add_bullet(doc, "Cmim psikologjik: $29.99 jo $30, $47.97 jo $48", "Psikologji: ")

add_heading_styled(doc, "Funneli i Shitjes", 2)
funnel = [
    ("Awareness", "Facebook/TikTok Ads, Pinterest", "CPM $5-15"),
    ("Interest", "Landing page, product video", "CTR target: 2-4%"),
    ("Desire", "Reviews, comparisons, urgency", "Add-to-cart rate: 8-15%"),
    ("Action", "Checkout optimized, trust badges", "Conversion rate: 2-4%"),
    ("Retention", "Email follow-up, upsell", "Repeat purchase: 15-25%"),
]
add_table(doc, ["Faza", "Taktikat", "KPI Target"], funnel)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. FINANCIAL MODEL
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "7. MODELI FINANCIAR / FINANCIAL MODEL", 1)

add_heading_styled(doc, "Kostot Fillestare (One-Time)", 2)
startup_costs = [
    ("Shopify Plan (Basic)", "$39/muaj", "Platforma e dyqanit"),
    ("Domen (.com)", "$15/vit", "neatnestco.com"),
    ("VPS Server (Lightsail)", "$5/muaj", "Bot intelligence system"),
    ("Proxy (residential)", "$15-30/muaj", "Per eBay/AliExpress access"),
    ("Anthropic API", "$5-10/muaj", "AI product analysis"),
    ("Logo + Branding", "$0-50", "Canva ose Fiverr"),
    ("Produkte test (samples)", "$50-100", "Per foto dhe testim"),
]
add_table(doc, ["Zeri", "Kosto", "Shenim"], startup_costs)
doc.add_paragraph("TOTAL FILLIMI: $120-250 (investim minimal)")

add_heading_styled(doc, "Projeksionet Mujore", 2)

projections = [
    ("Muaji 1", "5-10", "$30", "2%", "$150-600", "$50-200", "$100-400",
     "Testim, 5-10 produkte, $5/dite ads"),
    ("Muaji 2", "10-20", "$32", "2.5%", "$500-1,500", "$150-450", "$350-1,050",
     "Optimizim, scale winners"),
    ("Muaji 3", "20-40", "$35", "3%", "$1,500-4,000", "$450-1,200", "$1,050-2,800",
     "Scale agresive, 20+ produkte"),
    ("Muaji 6", "50-100", "$35", "3.5%", "$4,000-10,000", "$1,200-3,000", "$2,800-7,000",
     "Multi-niche, retargeting"),
    ("Muaji 12", "100-200", "$38", "4%", "$10,000-25,000", "$3,000-7,500", "$7,000-17,500",
     "Brand i njohur, LTV e larte"),
]
add_table(doc,
    ["Periudha", "Porosi/muaj", "AOV", "Conv%", "Revenue", "Kosto Ads", "Fitimi Bruto", "Shenim"],
    projections
)

add_heading_styled(doc, "Modeli i Fitimit per Produkt", 2)
product_model = [
    ("Cmimi Retail", "$35.00", "100%", "Cmimi ne dyqan"),
    ("Cmimi Burim (AliExpress)", "-$8.00", "-23%", "Kosto e produktit"),
    ("Transport", "-$3.00", "-9%", "ePacket/standard"),
    ("Platforma Shopify (2.9%+$0.30)", "-$1.32", "-4%", "Payment processing"),
    ("Tarifa Platformes (~15%)", "-$5.25", "-15%", "Apps, theme, overhead"),
    ("Reklama (~30%)", "-$10.50", "-30%", "Facebook/TikTok/Google"),
    ("FITIMI REAL", "$6.93", "19.8%", "Per njesi te shitur"),
]
add_table(doc, ["Zeri", "Shuma", "% e Retail", "Shenim"], product_model)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run(
    "Me 50 porosi/muaj x $6.93 fitim = $346.50 fitim mujor neto. "
    "Me 200 porosi/muaj = $1,386 fitim mujor neto. "
    "Objektivit: 200+ porosi/muaj brenda 6 muajve."
)
run.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. OPERATIONS PLAN
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "8. PLANI OPERACIONAL / OPERATIONS PLAN", 1)

add_heading_styled(doc, "Sistemi i Automatizuar (aktual)", 2)
doc.add_paragraph(
    "Sistemi i boteve ekzekutohet automatikisht cdo 24 ore nga VPS-ja. "
    "Raportet dergohen ne Telegram ne dy gjuhe (EN + SQ)."
)

ops_flow = [
    ("06:00", "Trend Scanner", "Skanon Google Trends per produkte ne rritje"),
    ("06:05", "Amazon Tracker", "Skanon 259+ produkte BSR ne 10 kategori"),
    ("06:15", "AliExpress Scanner", "Gjen furnizues me cmime te uleta"),
    ("06:20", "eBay Scanner", "Validon kerkesen me shitje reale"),
    ("06:25", "Price Monitor", "Kontrollon ndryshimet e cmimeve"),
    ("06:30", "Competitor Tracker", "Monitoron dyqanet konkurruese"),
    ("06:35", "Intelligence Engine", "Kryqezon te dhenat, scoring multi-sinjal"),
    ("06:40", "Report Generator", "Gjeneron HTML raporte + dergon ne Telegram"),
    ("06:45", "AI Analyzer", "Analiza e thelle me Claude AI (kur aktivizohet)"),
]
add_table(doc, ["Ora", "Boti", "Veprimi"], ops_flow)

add_heading_styled(doc, "Procesi i Porosise", 2)
order_flow = [
    ("1", "Klienti porosit ne Shopify", "Automatik"),
    ("2", "Njoftim ne email/Telegram", "Automatik"),
    ("3", "Porosi tek furnizuesi (AliExpress/CJ)", "Manual (5 min)"),
    ("4", "Furnizuesi dergon direkt tek klienti", "Automatik"),
    ("5", "Tracking number ne Shopify", "Manual (2 min)"),
    ("6", "Email follow-up pas 7 ditesh", "Automatik"),
]
add_table(doc, ["Hapi", "Veprimi", "Menyra"], order_flow)

add_heading_styled(doc, "Platformat e Furnizimit", 2)
suppliers = [
    ("AliExpress", "Primar", "Miliona produkte, cmime te uleta, ePacket 7-15 dite"),
    ("CJ Dropshipping", "Sekondar", "Magazine US, 3-7 dite transport, paketiim custom"),
    ("Zendrop", "Alternativ", "US warehouse, branding, 2-5 dite transport"),
    ("Spocket", "Premium", "Furnizues US/EU, 2-7 dite, cilesi me e larte"),
]
add_table(doc, ["Platforma", "Roli", "Avantazhet"], suppliers)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. ACTION PLAN 30-60-90 DAYS
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "9. PLANI I VEPRIMIT 30-60-90 DITE / ACTION PLAN", 1)

add_heading_styled(doc, "JAVA 1: Infrastruktura (Dite 1-7)", 2)
week1 = [
    "Aktivizo te gjitha 7 burimet e boteve (shto PROXY_URL per eBay/AliExpress)",
    "Shto $5 kredite Anthropic API per AI Analyzer",
    "Konfiguro cron schedule per ekzekutim automatik cdo 24 ore",
    "Dizajno logon e Neatnestco (Canva)",
    "Konfiguro Shopify theme profesionale",
    "Vendos politikat: Shipping, Returns, Privacy, Terms",
    "Konfiguro payment gateway (Shopify Payments / Stripe)",
]
for item in week1:
    add_bullet(doc, item)

add_heading_styled(doc, "JAVA 2-3: Produktet e Para (Dite 8-21)", 2)
week2 = [
    "Zgjidh 10 produkte nga kandidatet e botit (shiko seksionin 4)",
    "Porosit samples per 3-5 produktet kryesore",
    "Krijo listings profesionale: foto, pershkrim, SEO title",
    "Vendos cmimet sipas modelit te fitimit (seksioni 7)",
    "Konfiguro email marketing (Klaviyo free tier)",
    "Krijo faqen 'About Us' - historia e Neatnestco",
    "Testo procesin e porosise me nje test order",
]
for item in week2:
    add_bullet(doc, item)

add_heading_styled(doc, "JAVA 4: Lancimi (Dite 22-30)", 2)
week4 = [
    "Lanco Facebook Ads: $5/dite per 3 produkte kryesore",
    "Krijo TikTok business account + 3 video produktesh",
    "Vendos Facebook Pixel ne Shopify per tracking",
    "Lanco Google Shopping Ads per 5 produkte",
    "Fillo postimin ne Pinterest (5 pins/dite)",
    "Monitoro ROAS ditor nga Telegram reports",
    "Optimizo ads bazuar ne te dhenat e para",
]
for item in week4:
    add_bullet(doc, item)

add_heading_styled(doc, "MUAJI 2: Optimizim (Dite 31-60)", 2)
month2 = [
    "Analizo te dhenat: cili produkt shet me mire?",
    "Shkalo buxhetin e ads per produkte fituese (ROAS > 2x)",
    "Hiq produktet qe nuk shesin pas 500+ impressions",
    "Shto 10 produkte te reja nga raportet e fundit te botit",
    "Krijo lookalike audiences nga blereset e pare",
    "Fillo retargeting campaigns",
    "Shto produkte upsell/cross-sell",
    "Kerko micro-influencers (Instagram/TikTok): $50-200 per post",
]
for item in month2:
    add_bullet(doc, item)

add_heading_styled(doc, "MUAJI 3: Shkallezim (Dite 61-90)", 2)
month3 = [
    "Shkalo ne 20-40 produkte aktive",
    "Rrit buxhetin e ads ne $20-50/dite",
    "Shto nishe te re (Pet Products ose Automotive)",
    "Apliko per CJ Dropshipping per transport me te shpejte",
    "Fillo email sequences per klientet ekzistues",
    "Analizo LTV (Lifetime Value) te klienteve",
    "Krijo bundle offers per AOV me te larte",
    "Target: 40+ porosi/muaj, $1,000+ revenue",
]
for item in month3:
    add_bullet(doc, item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 10. RISKS & MITIGATION
# ═══════════════════════════════════════════════════════════════

add_heading_styled(doc, "10. RISQET DHE ZBUTJA / RISKS & MITIGATION", 1)

risks = [
    ("Koha e transportit", "E LARTE",
     "7-15 dite nga Kina",
     "Perdor CJ/Spocket me magazine US (3-5 dite). Komuniko qarte kohen e dergeses."),
    ("Konkurrenca", "MESATARE",
     "Shume dropshipper ne treg",
     "Diferencohu me brand, customer service, dhe selektim inteligjent te produkteve me bote."),
    ("Bllokimi i IP (bote)", "E ULET",
     "Amazon/eBay bllokojne IP",
     "Proxy residential ($15-30/muaj). Rotacion IP. Rate limiting."),
    ("Cilesia e produktit", "MESATARE",
     "Furnizuesi dergon produkt te dobet",
     "Porosit samples. Perdor furnizues me rating 4.5+. Return policy."),
    ("Ndryshimi i cmimeve", "E ULET",
     "Furnizuesi rrit cmimin",
     "Price Monitor bot gjurmon 24/7, alerta ne Telegram per ndryshime."),
    ("Facebook Ads ban", "MESATARE",
     "Account disabled",
     "Respekto politikat. Backup account. Diversifiko ne TikTok/Google."),
    ("Dispute/Chargeback", "E ULET",
     "Klienti kerkone para mbrapsht",
     "Tracking per cdo porosi. Responive customer service. Refund policy."),
]
add_table(doc, ["Risku", "Niveli", "Pershkrim", "Zbutja"], risks)

# ═══════════════════════════════════════════════════════════════
# CLOSING
# ═══════════════════════════════════════════════════════════════

doc.add_page_break()

add_heading_styled(doc, "PERFUNDIM / CONCLUSION", 1)

doc.add_paragraph(
    "Neatnestco ka te gjitha elementet per nje biznes te suksesshem dropshipping:"
)
add_bullet(doc, "Sistem inteligjence 24/7 me 10 bote te automatizuara", "Teknologji: ")
add_bullet(doc, f"{len(candidates)} kandidate te identifikuara, {len(cross_refs)} me shumefisha sinjale", "Te Dhena: ")
add_bullet(doc, "Shopify i konfiguruar, payment ready", "Platforme: ")
add_bullet(doc, "Raporte bilinguale (EN+SQ) ne Telegram cdo dite", "Raportim: ")
add_bullet(doc, "Formula e fitimit real qe llogarit TE GJITHA kostot", "Financiar: ")

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run(
    "HAPI I RRADHES: Aktivizo te gjitha burimet e boteve, zgjidh 10 produktet e para, "
    "dhe lanço reklamat brenda 7 ditesh. Objektivit: porositë e para brenda 30 ditesh."
)
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x00, 0xB8, 0x94)

doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— Signal Pilot v4 | Neatnestco Intelligence Engine —")
run.font.color.rgb = RGBColor(0x63, 0x6E, 0x72)
run.italic = True

# ── Save ──────────────────────────────────────────────────────

output = Path("reports") / f"Neatnestco_Business_Plan_{today}.docx"
doc.save(str(output))
print(f"Business plan saved: {output}")
print(f"Size: {output.stat().st_size:,} bytes")
